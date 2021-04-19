# -*- coding: cp1252 -*-
import os
import re
import datetime
import xlrd
from xlrd import open_workbook
import xlwt
import Tkinter
import tkFileDialog
import winsound
from operator import itemgetter
import openpyxl
import time
from openpyxl import load_workbook
import xlutils
from xlutils.copy import copy
from time import sleep
import sys
import ttk
import tkMessageBox
from Tkinter import *
#from tkinter import tkMessageBox
from PIL import ImageTk, Image
import webbrowser
import time
import Tkinter as tk
import Tkinter as tk
from Tkinter import *
from time import sleep
import ttk
import docx
from docx import Document


teams = range(100)
#####################################################################

FileTypeList_withoutResult = ['.java', '.JAVA', '.BAT', '.bat', '.ICE', '.ice']
ResultFileTypeList = ['.RES', '.res']
ExcelFiles=[".xlsx",".XLSX",".xls",".XLS"]
docx_type=['.docx', '.DOCX']

####################################String_Matching############################################
def BrowseFolderPath():
        print '***Select the Path for Synergy Data***'
        root = Tkinter.Tk()
        root.withdraw() #use to hide tkinter window
        currdir = os.getcwd()
        SynPath = tkFileDialog.askdirectory(parent=root, initialdir=currdir, title='Please Select the Path for Synergy Data')
        if len(SynPath) > 0:
                print "You have chosen: %s" % SynPath
        return SynPath
################################################################################
# Function to return the list of file path
################################################################################
def GetListOfFilesFromFolder(FolderPath, FileTypeList):
        try:
                FileNameList = []
                for FileType in FileTypeList:
                        for root, dirs, files in os.walk(FolderPath):
                                for f in files:
                                        if 'Support_Files' not in root and 'Archive' not in root:
                                                if f.endswith(FileType):
                                                        FileNameList.append(os.path.join(root, f))
        except:
                error_log_file.write('General Error\t[001]File name collection error!\n')
        return FileNameList

################################################################################
def GetSynVer_java(filename):
    f=open(filename,"r")
    syn_flag=0
    syn_version=""
    for line in f:
        if (("SYNERGY" in line.upper() and "INFORMATION" in line.upper()) or "HISTORY DATA" in line.upper()):
            syn_flag=1
        if("%VERSION:" in line.upper() and syn_flag==1):
            syn_version=line.replace("%version","").replace("#","").replace("\n","").replace("\r","")\
                         .replace("%","").replace(" ","\t").replace("\t","").replace(":","").replace("**","")
            syn_flag=2
    try:
            return  float(syn_version)
    except:
            pass
################################################################################





################################################################################
def GetSynVer_bat(filename):
    f=open(filename,"r")
    syn_flag=0
    syn_version=""
    for line in f:
        if (("SYNERGY" in line.upper() and "INFORMATION" in line.upper()) or "HISTORY DATA" in line.upper()):
            syn_flag=1
        if("%VERSION:" in line.upper() and syn_flag==1):
            syn_version=line.replace("%version","").replace("#","").replace("\n","").replace("\r","")\
                         .replace("%","").replace(" ","\t").replace("\t","").replace(":","")
            syn_flag=2
    try:
            return  float(syn_version)
    except:
            pass
################################################################################





################################################################################
def GetSynVer_ice(filename):
    f=open(filename,"r").read().splitlines();
    syn_line=[]
    syn_version=""
    syn_flag=0;
    for line in f:
            if (("SYNERGY" in line.upper() and "INFORMATION" in line.upper()) or "HISTORY DATA" in line.upper()):
                    syn_flag=1
            if("#** %VER" in line.upper() or "#**  %VER" in line.upper() and syn_flag==1):
                    syn_line=line.split(':')
                    syn_version=syn_line[1].replace("\t","").replace("%","").replace("\n","").replace(" ","")\
                                 .replace("created_by","")


    try:
            return  float(syn_version)
    except:
            pass
################################################################################





################################################################################
def GetManVer_java(filename):
        f=open(filename,"r")
        man_version=""
        for line in f:
                if "AUTHOR(S):" in line.upper():
                        man_version=line.replace("**      Author(s):  Bite2SiteTranslator.java","").replace("(","").replace("AUTOGENERATED","")\
                                     .replace(")","").replace("V","").replace(" ","").replace("\t","").replace("\n","")
        try:
                return float(man_version)
        except:
                pass

#################################################################################################################################


################################################################################
def GetManVer_bat(filename):
        f=open(filename,"r")
        man_verflag=0
        man_version=""
        countn=0
        HistoryList=[]
        HistoryLastLineList=[]
        verlist=[]
        for line in f:
                if " VER." in line.upper() or " VER " in line.upper() or "\tVER\t" in line.upper() \
                   or " VER\t" in line.upper() or "\tVER " in line.upper():
                        man_verflag=1

                if (("SYNERGY" in line.upper() and "INFORMATION" in line.upper()) or "HISTORY DATA" in line.upper()):
                        man_verflag=man_verflag+1

                if(man_verflag==1):
                        i=0

                        for entry in line.split(" "):
                                if(i<7):
                                        try:
                                                verlist.append(float(entry))
                                        except:
                                                pass
                                i=i+1
        try:
                return max(verlist)
        except:
                pass

#################################################################################################################################


################################################################################
def GetManVer_ice(filename):
        f=open(filename,"r").read().splitlines();
        man_verflag=0
        author_flag=0
        global type_flag
        type_flag=0
        verlist=[]
        for line in f:
                if("AUTHOR(S):" in line.upper()):
                        author_flag=1
                if " VER." in line.upper() or " VER " in line.upper() or "\tVER\t" in line.upper() \
                   or " VER\t" in line.upper() or "\tVER " in line.upper():
                        man_verflag=1

                if (("SYNERGY" in line.upper() and "INFORMATION" in line.upper()) or "HISTORY DATA" in line.upper() or "HISTORY" in line.upper()):
                        man_verflag=2
                if (author_flag==1 and man_verflag==2):
                        if("#** %VER" in line.upper() or "#**  %VER" in line.upper()):
                                verlist=line.split(':')
                                verlist[1]=verlist[1].replace("\t","").replace("%","").replace("\n","").replace(" ","")\
                                             .replace("created_by","")
                                type_flag=1

                if(man_verflag==1):
                        type_flag=2
                        i=0

                        for entry in line.split(" "):
                                if(i<7):
                                        try:
                                                verlist.append(float(entry))
                                        except:
                                                pass
                                i=i+1
        try:
                if(type_flag==1):
                        return float(verlist[1])
                else:
                        return max(verlist)
        except:
                pass

#################################################################################################################################



#################################################################################################################################
def GetAuthName_java(filename,ManualVer):
        return "Bite2SiteTranslator.java"





#################################################################################################################################

#################################################################################################################################
def GetAuthName_bat(filename,ManualVer):
        f=open(filename,"r")
        auth_flag=0
        author_names=[]

        for line in f:
                if " VER." in line.upper() or " VER " in line.upper() or "\tVER\t" in line.upper() \
                   or " VER\t" in line.upper() or "\tVER " in line.upper():
                        auth_flag=1

                if (("SYNERGY" in line.upper() and "INFORMATION" in line.upper()) or "HISTORY DATA" in line.upper()):
                        auth_flag=auth_flag+1
                if(auth_flag==1):
                        BlankLineflag=0
                        for i in range(0,8):
                                try:

                                        if(float(line[i])):
                                                BlankLineflag=1
                                                break
                                except:
                                        pass

                        if(BlankLineflag==1):
                                line=line.split("  ")
                                templines=[]
                                for entry in line:
                                        if(entry!=""):
                                                if(entry.upper()=="HISTORY"):
                                                        entry=entry.upper().replace("HISTORY","")
                                                entryWithoutnum=''.join([i for i in entry if((i >= 'a' and i <= 'z') or (i >= 'A' and i <= 'Z') or i==" " or i=="\t" or i==".")])
                                                templines.append(entryWithoutnum)




                                try:

                                           if(len(str(templines[2]))>2):
                                                   author_names.append(templines[2])
                                           else:
                                                  author_names.append(templines[1])



                                except:
                                        pass

        try:

                if(author_names[-1]=="" or len(str(author_names[-1]))<=2):

                        return str(author_names[-2])
                else:

                        return str(author_names[-1])
        except:
                pass
#################################################################################################################################


#################################################################################################################################
def GetAuthName_ice(filename,ManualVer):
        f=open(filename,"r").read().splitlines();
        auth_flag=0
        author_names=[]

        for line in f:
                if " VER." in line.upper() or " VER " in line.upper() or "\tVER\t" in line.upper() \
                   or " VER\t" in line.upper() or "\tVER " in line.upper():
                        auth_flag=1

                if (("SYNERGY" in line.upper() and "INFORMATION" in line.upper()) or "HISTORY DATA" in line.upper()or "HISTORY" in line.upper()):
                        auth_flag=2
                if(auth_flag==1):
                        BlankLineflag=0
                        for i in range(0,8):
                                try:

                                        if((line[i])):
                                                BlankLineflag=1
                                                break

                                except:
                                        pass

                        if(BlankLineflag==1):
                                line=line.split("  ")
                                templines=[]
                                for entry in line:
                                        if(entry!=""):
                                                if(entry.upper()=="HISTORY"):
                                                        entry=entry.upper().replace("HISTORY","")
                                                entryWithoutnum=''.join([i for i in entry if((i >= 'a' and i <= 'z') or (i >= 'A' and i <= 'Z') or i==" " or i=="\t" or i==".")])
                                                templines.append(entryWithoutnum)




                                try:

                                           if(len(str(templines[2]))>2):
                                                   author_names.append(templines[2])
                                           else:
                                                  author_names.append(templines[1])



                                except:
                                        pass

        try:

                if(author_names[-1]=="" or len(str(author_names[-1]))<=2):

                        return str(author_names[-2])
                else:

                        return str(author_names[-1])
        except:
                pass
#################################################################################################################################

#################################################################################################################################
def GetAuthName_ice_t1(filename):
    Auth_name=""
    f=open(filename,"r").read().splitlines();
    for line in f:
        if("AUTHOR(S):" in line.upper()):
            Auth_name=line.replace("#**  Author(s): ","")
    try:
            return Auth_name
    except:
            pass


#################################################################################################################################

#################################################################################################################################
def GetResult_java(filename):

        res=""
        f=open(filename,"r")
        for line in f:
                if('OVERALL TEST SCRIPT STATUS' in line.upper()):
                        res=line.replace("Overall Test Script Status","").replace("=>","").replace(" ","").replace("\t","").replace("\n","").replace("ED","")
        return res.title()






#################################################################################################################################


#################################################################################################################################
def GetResult_bat(filename):
        failed=1
        res=""
        f=open(filename,"r")
        for line in f:
            if('#TEST FAILED#' in line.upper()):
                failed=0
                break
        if failed==0:
            res= 'Fail'
        else:
            res= 'Pass'
        return res






#################################################################################################################################

#################################################################################################################################
def GetResult_ice(filename):
        res=""
        f=open(filename,"r").read().splitlines();
        for line in f:
                if('$8' in line.upper() or '$5' in line.upper() or '$4' in line.upper()):
                        res=line.replace("$5","").replace("=","").replace(" ","").replace("\"","").replace("\n","").replace("$4","").replace("$8","")
        return res.title()


#################################################################################################################################

#################################################################################################################################
def GetSVCPNo(filename):
        f=open(filename,"r").read().splitlines();
        SVCPlist=[]
        for line in f:
               if "Test Case Identifier" in line:
                       line=line.replace("Test Case Identifier","").replace("#","").replace(":","").replace("\t"," ")\
                       .replace(" ","").replace("*","")
                       SVCPlist.append(line)
        return SVCPlist

#################################################################################################################################

#################################################################################################################################
def GetTPSID(entry):
        try:
                svcp_col=0
                tps_col=0
                for col in range(1,sheet_mainapp.max_columns+1):
                        if("ID" in str(sheet_mainapp.cell(row=1,columns=col).value)):
                                tps_col=col

                        if("Out-links at depth 1" in str(sheet_mainapp.cell(row=1,columns=col).value)):
                                svcp_col=col


                if("," in entry):
                        entry_list=entry.split(",")
                        tpsList=[]
                        for entry in entry_list:
                                for row in range(1,sheet_mainapp.max_row+1):
                                        if(str(entry).replace("\n","").replace("\t","").replace("\r\n","") in "".join(sheet_mainapp.cell_value(row,svcp_col))):
                                                tpsList.append(sheet_mainapp.cell_value(row,tps_col)+"\n")
                        try:
                                return (tpsList)
                        except:
                                pass
                else:
                        for row in range(1,sheet_mainapp.max_row+1):
                                if(str(entry).replace("\n","").replace("\t","").replace("\r\n","") in "".join(sheet_mainapp.cell_value(row,svcp_col))):
                                        return (sheet_mainapp.cell_value(row,tps_col))
        except:

                error_log_file.write("TPS_SVCP_TR.xlsx is missing in the synergy folder\n")

                tkMessageBox.showinfo("Error","TPS_SVCP_TR.xlsx is missing in the synergy folder")
                window.destroy()

                sleep(10)
                sys.exit()
#################################################################################################################################

#################################################################################################################################

def GetTCid_bat(filename):
        TC_ID=""
        flag=0
        temp=""
        f=open(filename,"r").read().splitlines();
        for line in f:
                if("TEST CASE ID" in line.upper()):
                        flag=1
                        temp=line.replace("Test Case ID","").replace(":","").replace(" ","").replace("\n","")
                        TC_ID=TC_ID+temp+" \r\n"
                elif("END OF TEST CASE" in line.upper() and flag==0):
                        temp=line.replace("End of ","").replace("-","").replace("\n","")
                        TC_ID=TC_ID+temp+" \r\n"


        return TC_ID

#################################################################################################################################

#################################################################################################################################

def GetTCid_java(filename):
        TC_ID=""
        tcId=[]
        temp=[]
        f=open(filename,"r").read().splitlines();
        for line in f:
                if("TC" in line.upper() and "FAULT" in line.upper()):
                        temp=line.split(":")
                        if(temp[0] not in tcId):
                                TC_ID=TC_ID+temp[0]+" \r\n"
                        tcId.append(temp[0])


        return TC_ID

#################################################################################################################################

#################################################################################################################################


def GetTCid_ice(filename):
        f=open(filename,"r").read().splitlines();
        tcid=""
        for line in f:
            if("#**  TEST PROCEDURE NUMBER" in line.upper()):

                tcid=line.replace("Test Procedure Number","").replace("\t","").replace(":","").replace(" ","").replace("\n","").replace("#","").replace("*","")
                break
            if("#**  TEST CASE" in line.upper() and "BASELINE"  not in line.upper()):
                tcid=line.replace("Test Case","").replace("\t","").replace(":","").replace(" ","").replace("\n","").replace("#","").replace("*","")
                break
        return tcid



#################################################################################################################################

#################################################################################################################################
def GetSWRDandSWDD(entry):

        try:
                svcp_col=0
                swrd_col=0
                for col in range(1,sheet_mainapp_1.max_column+1):
                        if("ID" in str(sheet_mainapp_1.cell(row=0,column=col).value)):
                                svcp_col=col
                        if("Out-links at depth 1" in str(sheet_mainapp_1.cell(row=0,column=col).value)):
                                swrd_col=col

                mlct_col=0
                swrd_mlct_col=0
                for col in range(1,sheet_mainapp_2.max_column+1):
                        if("Out-links at depth 1" in str(sheet_mainapp_2.cell(row=0,column=col).value)):
                                mlct_col=col
                        if("Out-links at depth 2" in str(sheet_mainapp_2.cell(row=0,column=col).value)):
                                swrd_mlct_col=col

                if("," in entry):
                        entry_list=entry.split(",")
                        swrdandswddList=[]
                        for entry in entry_list:
                                for row in range(1,sheet_mainapp_1.max_row+1):
                                        if(str(entry).replace("\n","").replace("\t","").replace("\r\n","") in "".join(sheet_mainapp_1.cell_value(row,svcp_col))):
                                                swrdandswddList.append(sheet_mainapp_1.cell_value(row,swrd_col)+"\n")
                                for row in range(1,sheet_mainapp_2.max_row+1):
                                        if(str(entry).replace("\n","").replace("\t","").replace("\r\n","") in "".join(sheet_mainapp_2.cell_value(row,mlct_col))):
                                                swrdandswddList.append(sheet_mainapp_2.cell_value(row,swrd_mlct_col)+"\n")
                        return (swrdandswddList)
                else:
                        for row in range(1,sheet_mainapp_1.max_row+1):
                                if(str(entry).replace("\n","").replace("\t","").replace("\r\n","") in "".join(sheet_mainapp_1.cell_value(row,svcp_col))):
                                        return (sheet_mainapp_1.cell_value(row,swrd_col))
                        for row in range(1,sheet_mainapp_2.max_row+1):
                                if(str(entry).replace("\n","").replace("\t","").replace("\r\n","") in "".join(sheet_mainapp_2.cell_value(row,mlct_col))):
                                        return (sheet_mainapp_2.cell_value(row,swrd_mlct_col))
        except:
                error_log_file.write("Main_app_SVCP_SWRD_SWDD.xlsx and\\or MCLT_SVCP_inlink_outlink.xlsx is missing in the synergy folder\n")
                tkMessageBox.showinfo("Error","Main_app_SVCP_SWRD_SWDD.xlsx and\\or MCLT_SVCP_inlink_outlink.xlsx is missing in the synergy folder")
                window.destroy()
                sleep(10)
                sys.exit()
#################################################################################################################################
def GetSVCPList():
        try:
                svcp_col=0
                swrd_col=0
                for col in range(1,sheet_mainapp_1.max_column):
                        if("ID" in str(sheet_mainapp_1.cell(row=0,column=col).value)):
                                svcp_col=col
                        if("DS10793/327" in str(sheet_mainapp_1.cell(row=0,column=col).value)):
                                swrd_col=col

                for i in range(sheet_mainapp_1.max_row):
                       if "Objective:" in str(sheet_mainapp_1.cell(i, swrd_col).value) or "Objective_" in str(sheet_mainapp_1.cell(i, swrd_col).value) or "To verify" in str(sheet_mainapp_1.cell(i, swrd_col).value)\
                          or "To Verify" in sheet_mainapp_1.cell_value(i, swrd_col):
                                  SVCP_List.append((sheet_mainapp_1.cell(i, svcp_col).value))
        except:
                error_log_file.write("Main_app_SVCP_SWRD_SWDD.xlsx is missing in the synergy folder\n")
                tkMessageBox.showinfo("Error","Main_app_SVCP_SWRD_SWDD.xlsx is missing in the synergy folde")
                window.destroy()
                sleep(10)
                sys.exit()

#################################################################################################################################

def GetResTestScript(filename):
        listts=""
        f=open(filename,"r")
        for line in f:
                if("TEST FILE NAME" in line.upper()):
                        listts=line.replace("Test File Name","").replace("=>","").replace("\n","").replace("\t","").replace(" ","").replace("*","").replace("TestfileName:","")
                        break
                if("SCRIPT FILE NAME" in line.upper()):
                        listts=line.replace("Script File Name","").replace(":","").replace("\n","").replace("\t","").replace(" ","").replace("TestfileName:","")
                        break
                if("*** TEST FILE NAME" in line.upper()):
                        listts=line.replace("*** Test file Name","").replace(":","").replace("\n","").replace("\t","").replace(" ","").replace("TestfileName:","")
                        break
                if("#**               * " in line.upper() or "#**          * " in line.upper() or "#**       * " in line.upper()\
                   or "#**             * " in line.upper() or "#**           * " in line.upper() or "#**                   * " in line.upper()\
                   or "#**               *" in line.upper() and ("***" not in line.upper()) or "#**                             *S" in line.upper()):
                        listts=line.replace("#","").replace("*","").replace("\n","").replace("\t","").replace(" ","").replace("TestfileName:","")
                        break
                if("SCRIPT FILE NAME" in line.upper()):
                        listts=line.replace("Script File Name","").replace(":","").replace("\n","").replace("\t","").replace(" ","").replace("TestfileName:","")
                        break
                if("*** TEST FILE NAME" in line.upper()):
                        listts=line.replace("*** Test file Name","").replace(":","").replace("\n","").replace("\t","").replace(" ","").replace("TestfileName:","")
                        break
                if("  TEST FILE NAME    =>" in line.upper()):
                        listts=line.replace("  Test File Name    =>","").replace("\n","").replace("\t","").replace(" ","").replace("TestfileName:","")

        return listts
#################################################################################################################################


#################################################################################################################################
def GetResultTestScriptVersion(filename):
    listsv=""
    f=open(filename,"r")


    for line in f:

        if("*** TEST FILE VERSION" in line.upper()):
                listsv=line.replace("*** Test file version","").replace("\t","").replace(":","").replace(" ","").replace("\n","")
                break
        if("Script Version" in line.upper()):
                listsv=line.replace("Script Version","").replace("\t","").replace(":","").replace(" ","").replace("\n","")
                break
        if("TEST FILE VERSION" in line.upper()):
                listsv=line.replace("Test File Version","").replace("\t","").replace("=>","").replace(" ","").replace("\n","")
                break

    try:
            return  float(listsv)
    except:
            pass



#################################################################################################################################


#************************************************************
#* Procedure Name:          String_Matching
#* Procedure Description:
#*                          finding the string matching count
#*                          in file
#* Input Parameters:
#*                          Filename
#*                          pattern
#* Output Parameters:
#*                          Number of string count.
#************************************************************
def String_Matching(filename,pattern):
    Count = 0
##    with (file) as f:
    f=open(filename,"r").read().splitlines();
    for x in f:
        if pattern in x:
                Count =  Count + 1

    return Count

def Reverse_reading(file,pattern):

    ##    with (file) as f:
    f = open(file, "r").read().splitlines();
    for x in reversed(f):
        if pattern in x:
            x = x.split(":")
            x = x[0].replace("TC","")
            print x
            x = int(x)
            break

    return x



def GetScriptName(res_name):
        script_name=""
        if(res_name[-6:]=="_SPDA1" or res_name[-6:]=="_SPDA2"):
                res_name=res_name[:-6]
        res_name=res_name+"."
        for file_path in MasterFileNameList:
                if(res_name in file_path):
                        script_name=file_path
        return script_name


def Docx_GetScriptName(res_name):
        script_name=""
        Res_Name_1=res_name.replace("_RESULT","").\
                       replace("_Result","").replace("_result","").\
                       replace("..",".").replace(" ",".").replace("..",".")
        for file_path in DocxFileList:
                if(Res_Name_1.upper() in file_path.upper()):
                         script_name=file_path
        return script_name





def month_string_to_number(string):
    m = {
        'jan': 1,
        'feb': 2,
        'mar': 3,
        'apr':4,
         'may':5,
         'jun':6,
         'jul':7,
         'aug':8,
         'sep':9,
         'oct':10,
         'nov':11,
         'dec':12
        }
    s = string.strip()[:3].lower()
    try:
        out = m[s]
        return out
    except:
        print(string +" Is not a month")

def Time_Stamp(script,result):
 S_Date_Time=""
 R_Date_Time=""
 syn_Res=[]
 syn_Script=[]
 Time_Res=[]
 Time_Script=[]
 Script_file_path = open(script,'r')
 Result_file_path = open(result,'r')
 if(script.split(".")[-1] == "bat" ) or (script.split(".")[-1] == "ice" ) or (script.split(".")[-1] == "java" ) or (script.split(".")[-1] == "c" ):
  for line in Script_file_path:
    if ("modify" in line.lower()):
      S_Date_Time=line.replace("%modify_time:","").replace("%","").replace("#","").replace("*","").strip()
      S_Date_Time =  S_Date_Time.replace ("  "," ")
      syn_Script= S_Date_Time.split(" ")
 if(result.split(".")[-1] =="res"):
  for line in Result_file_path:
    if ("Test Date and Time" in line):
      R_Date_Time=line.replace("Test Date and Time","").replace("%","").replace(":"," ").replace("\n", "").strip()
      R_Date_Time = R_Date_Time.replace ("  "," ")
      syn_Res= R_Date_Time.split(" ")

  if len (syn_Res)< 9 or len(syn_Script)< 5:
    return "Time stamp is mismatching"

 if (syn_Res[8]==syn_Script[4]):
  if (month_string_to_number(syn_Res[1])==month_string_to_number(syn_Script[1])):

    if int (syn_Res[2])==int (syn_Script[2]):

       Time_Script =syn_Script[3].split(":")
       if 'PM' in syn_Res[6]:
          syn_Res[3] = int (syn_Res[3])+12
       if int(syn_Res[3]) == int (Time_Script[0]):
          if int(syn_Res[4]) == int (Time_Script[1]):
            if int(syn_Res[5]) >= int (Time_Script[2]):
              return "Time stamp is matching"
            else:
              return "Time stamp is mismatching"
          elif int(syn_Res[4]) > int (Time_Script[1]):
            return "Time stamp is matching"
          else:
            return "Time stamp is mismatching"
       elif int(syn_Res[3]) > int (Time_Script[0]):
          return "Time stamp is matching"
       else:
          return "Time stamp is mismatching"
    elif int (syn_Res[2])>int (syn_Script[2]):
       return "Time stamp is matching"
    else:
      return "Time stamp is mismatching"
  elif (month_string_to_number(syn_Res[1])> month_string_to_number(syn_Script[1])):
    return "Time stamp is matching"
  else:
    return "Time stamp is mismatching"
 elif (syn_Res[8]>syn_Script[4]):
  return "Time stamp is matching"
 else:
  return "Time stamp is mismatching"

################################################################################

def Checklist(entry_list):
        start_time = time.time()

        flag=0

#************************************************************
#* Procedure Name:          Review checklist
#* Procedure Description:
#*                          finding the review checklist for TC/TS/TR
#*                          in file
#* Input Parameters:
#*                          TC:
#                           SVCP ID
#                           TS:
#                           Script name / Final version
#                           TR:
#                           Result name / Result version
#* Output Parameters:
#*                          Return.
#************************************************************/




# r=root, d=directories, f = files
        for r, d, f in os.walk(Synergy_path):
          if ("Archive" not in r):
            for file in f:

                flag=0
                if ".xlsm" in file and ("STS" in file or "STC" in file or "STR" in file):
                    temp = file
                    file = r+"\\"
                    file = file +temp   # concatinate the file path
                    try:
                        # Load one by one excel file
                        wb_obj = openpyxl.load_workbook(filename = file, read_only = True, keep_vba = False)
                        sheet_obj = wb_obj['Review Tracker']
                        #print file

                        row = 1 #initialize the row value to 1 for each excel file
                        SVCP_ID_CELL = sheet_obj.cell(row=row, column=2)
                    except:
                        flag=1
                        error_log_file.write("No Review Tracker tab in file: %s \\ File  IO error\n" % (file)) # FILE IO error / Review Tracker sheet not found
                    if(flag==0):

                        while SVCP_ID_CELL.value:   # read each row value /SVCP ID/SCRIPT name/RESULT name
                            # print i
                            SVCP_ID_CELL = sheet_obj.cell(row=row, column=2)
                            Final_Revision_ID_CELL = sheet_obj.cell(row=row, column=5)

                            Result_Version_ID_CELL = sheet_obj.cell(row=row, column=3)


                            for entry in entry_list:

                                    if entry[0] == SVCP_ID_CELL.value and Final_Revision == Final_Revision_ID_CELL.value:

                                            entry[3]= entry[3]+file.split("\\")[-1] + "  SVCP version: "+str(Final_Revision_ID_CELL.value)+"\n"
                                    elif entry[0] == SVCP_ID_CELL.value and Final_Revision != Final_Revision_ID_CELL.value:

                                            entry[3]= entry[3] + file.split("\\")[-1] + "  SVCP version is not matching: "+str(Final_Revision_ID_CELL.value)+"\n"


                                    if entry[1] == SVCP_ID_CELL.value and entry[6] == Final_Revision_ID_CELL.value:

                                            entry[4]= entry[4]+ file.split("\\")[-1] + "  Script version: "+str(Final_Revision_ID_CELL.value)+"\n"
                                    elif entry[1] == SVCP_ID_CELL.value and entry[6] != Final_Revision_ID_CELL.value:

                                            entry[4]= entry[4]+ file.split("\\")[-1] + "  Script version is not matching: "+str(Final_Revision_ID_CELL.value)+"\n"


                                    if entry[2] == SVCP_ID_CELL.value and entry[7] == Result_Version_ID_CELL.value:

                                            entry[5]= entry[5]+ file.split("\\")[-1] + "  Result version: "+str(Result_Version_ID_CELL.value)+"\n"
                                    elif entry[2] == SVCP_ID_CELL.value and entry[7] != Result_Version_ID_CELL.value:

                                            entry[5]= entry[5]+file.split("\\")[-1] + "  Result version is not matching: "+str(Result_Version_ID_CELL.value)+"\n"

                            row += 1 #increment the row value of column 2
                        wb_obj.close()


        for entry in entry_list:
                if(entry[3]==" "):
                        entry[3]="File Not Found"
                if(entry[4]==" "):
                        entry[4]="File Not Found"
                if(entry[5]==" "):
                        entry[5]="File Not Found"
        print("--- %s seconds ---" % (time.time() - start_time))

        return entry_list



def GetResult_version(filename):
        f=open(filename,"r")
        ver=0

        for line in f:

                if("%version" in line):
                        try:
                                ver=int(line.split(":")[-1].replace("%","").strip())
                        except:
                                i=0
                        break
        return ver


def RemoveDuplicate(duplicate):
    final_list = []
    for num in duplicate:
        if num not in final_list:
            final_list.append(num)
    return final_list


#################################################################################################################################
def Docx_GetSVCPNo(filename):
        document = Document(filename)
        SVCPlist=[]

        docx_svcp_flag=0
        for line in document.paragraphs:
               if "Test Case Identifier" in line.text:
                       line=line.text.replace("Test Case Identifier","").replace("#","").replace(":","").replace("\t"," ")\
                       .replace(" ","").replace("*","")
                       SVCPlist.append(line)
                       docx_svcp_flag=1
        if docx_svcp_flag==1:
                return SVCPlist, docx_svcp_flag
        else:
                SVCPlist=[" "]
                return SVCPlist, docx_svcp_flag

#################################################################################################################################


#################################################################################################################################

def Docx_GetSWRDandSWDD_1(filename):
        document = Document(filename)
        swrdandswddList=[]
        range_flag_begin=0
        range_flag_end=0
        lines=""
        i=0
        for line in document.paragraphs:
                if("Normal/Robustness test" in line.text or "Normal/Robust Testing" in line.text):
                        range_flag_end=1
                        break
                if( "REQUIREMENTS UNDER TEST" in line.text.upper()):
                        range_flag_begin=1
                        lines=line.text.replace("Requirements under Test","").replace("SWRD:","").replace("REQ","").replace(":","").replace(" ","").replace("*","")\
                               .replace("#Requirementsundertest","").replace("#","").replace("RequirementsUnderTest","")
                        swrdandswddList.append(lines+"  ")
                        continue
                if(range_flag_begin==1 and range_flag_end==0):
                        lines=line.text.replace("Requirements under Test","").replace("REQ","").replace("SWRD:","").replace("SWDD:","").replace(" ","").replace("*","")\
                               .replace("#Requirementsundertest","").replace("#","").replace(":","").replace("RequirementsUnderTest","")
                        swrdandswddList.append(lines+"  ")
        return swrdandswddList




#################################################################################################################################


#################################################################################################################################



def Docx_GetAuthor(filename):
        document = Document(filename)
        author=""
        lines=[]
        for line in document.paragraphs:
                        if('AUTHOR(S)' in line.text.upper()):

                                lines=line.text.split(":")
                                author=lines[1].replace("/"," ").replace("\\"," ").replace(","," ").replace("*","")
                                break
        return author



#################################################################################################################################


#################################################################################################################################



def Docx_GetResultVersion(filename):
        document = Document(filename)
        ver=""
        lines=[]
        for line in document.paragraphs:
                        if('TEST VERSION:' in line.text.upper()):

                                lines=line.text.split(":")
                                ver=lines[1].replace("\n"," ").replace("\r"," ").replace(" ","")
                                break
        return int(ver)



#################################################################################################################################

#################################################################################################################################



def Docx_GetResultTCcount(filename):
        document = Document(filename)
        count=0
        flag1=0
        flag2=0
        flag3=0
        for line in document.paragraphs:
                        if('TEST CASE:' in line.text.upper()):
                                flag1=1
                                count=count+1
        if(flag1==0):
                for line in document.paragraphs:
                        if('TEST PROCEDURE NUMBER' in line.text.upper()):
                                flag2=1
                                count=count+1
        
        if(flag1==0 and flag2==0):
                for line in document.paragraphs:
                        if('TEST CASE IDENTIFIER' in line.text.upper()):
                                flag1=1
                                count=count+1
                

                        
        return count



#################################################################################################################################


#################################################################################################################################



def Docx_GetResultTCID(filename):
        document = Document(filename)
        tcList=[]
        flag1=0
       
       
        for line in document.paragraphs:
                        if('TEST CASE:' in line.text.upper()):
                                flag1=1
                                tc=line.text.split(":")[1].replace("\r"," ").replace("\n"," ").replace("/","").replace("*","").replace("#","").replace(" ","")
                                tcList.append(tc+"\n")
        if(flag1==0):
                for line in document.paragraphs:
                        if('TEST PROCEDURE NUMBER' in line.text.upper()):
                                tc=line.text.split(":")[1].replace("\r"," ").replace("\n"," ").replace("/","").replace("*","").replace("#","").replace(" ","")
                                tcList.append(tc+"\n")
                

                        
        return tcList



#################################################################################################################################

#################################################################################################################################
def Docx_GetResult(filename):

        res=""
        document = Document(filename)
        tables = document.tables
        flag=0
        flag_1=0



        for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for line in cell.paragraphs:
                                if('OVERALL TEST STATUS PASSED/FAILED' in line.text.upper() or 'Overall Test Pass/Fail Status' in line.text):
                                        flag=1
                                if(('PASS' in line.text.upper() or 'FAIL' in line.text.upper()) and (flag==1 and 'OVERALL TEST STATUS PASSED/FAILED' not in line.text.upper())):
                                        res=line.text.upper().replace("ED","").replace("**    Overall Test Pass/Fail Status:","").replace("OVERALL TEST STATUS PASSED/FAILED","")
                                        flag_1=1

                                        break
        if(flag_1==0):
                for line in document.paragraphs:
                        if('OVERALL TEST STATUS PASSED/FAILED' in line.text.upper() or 'Overall Test Pass/Fail Status' in line.text):
                                
                                res=line.text.upper().replace("OVERALL TEST STATUS PASSED/FAILED","").replace("Overall Test Pass/Fail Status","").replace(":","").replace(" ","").replace("\t","")\
                                     .replace("*","").replace("#","").replace("\n","").replace("ED","").replace("**    Overall Test Pass/Fail Status:","").replace("OVERALLTESTPASS/FAILSTATUS","")




        return res.title()






#################################################################################################################################




################################################################################
def ConsolidateData(ExcelFileList):
    wb=xlwt.Workbook()
    sheet1=wb.add_sheet("Analysis")
 
    docx_svcp_flag=0
    result=""
    tc_id=""
    file_list=[]
    svcp_in_files=[]
    Entered_docx=0

    Test_script_case_count=0
    Test_Result_case_count=0
    now = datetime.datetime.now()
    date_time=""
    date_time=now.strftime("_%m-%d-%Y_%Hh-%Mm-%Ss")
    sheet1.write(0,0,"SCRIPT NAME")
    sheet1.write(0,1,"PATH")
    sheet1.write(0,2,"SCRIPT SYN VER")
    sheet1.write(0,3,"SCRIPT MANUAL VER")
    sheet1.write(0,4,"VERSION MATCH?")
    sheet1.write(0,5,"AUTHOR(S)")
    sheet1.write(0,6,"SVCP ID")
    sheet1.write(0,7,"TPS ID")
    sheet1.write(0,8,"SWRD/SWDD")
    sheet1.write(0,9,"RESULT FILE")
    sheet1.write(0,10,"PATH")
    sheet1.write(0,11,"Result Version")
    sheet1.write(0,12,"SCRIPT VERSION IN RESULT FILE")
    sheet1.write(0,13,"RESULT STATUS")
    sheet1.write(0,14,"TC ID")
    sheet1.write(0,15,"Test Script Case Count")
    sheet1.write(0,16,"Test Result Case Count")
    sheet1.write(0,17,"Test Case Count match?")
    sheet1.write(0,18,"Synergy version of script and result match?")
    sheet1.write(0,19,"Time Stamp Match?")
    sheet1.write(0,20,"Test Case Checklist")
    sheet1.write(0,21,"Test Script Checklist")
    sheet1.write(0,22,"Test Result Checklist")

    i=1




    for filename_res in ResultFileList:
            FileWithoutPath_res=filename_res.split("\\")[-1]
            if(FileWithoutPath_res.endswith(".docx")):
                    filename=Docx_GetScriptName(FileWithoutPath_res.replace(".docx",""))
            else:
                    filename=GetScriptName(FileWithoutPath_res.replace(".res",""))
            if(filename.endswith(tuple(FileTypeList_withoutResult)) or\
               filename.endswith(tuple(docx_type))):
                FileWithoutPath=filename.split("\\")[-1]
                ScriptSynVer=""
                ManualVer=""
                TPS=""
                SWRDandSWDD=""
                #Function calls for .java
                if(filename.endswith(".java")):
                        ScriptSynVer=GetSynVer_java(filename)
                        ManualVer=GetManVer_java(filename)
                        SVCPNo=[' ']
                        Test_script_case_count = String_Matching(filename,'Test Case:')
                        AuthName=GetAuthName_java(filename,ManualVer)
                #Function calls for .bat
                elif(filename.endswith(".bat")):
                        if(filename.split(FileWithoutPath)[0].find('\\SSIT') == -1):
                                ScriptSynVer=GetSynVer_bat(filename)
                                SVCPNo=GetSVCPNo(filename)
                                Test_script_case_count = String_Matching(filename,'Test Procedure Number')
                                ManualVer=GetManVer_bat(filename)
                                AuthName=[element+" " for element in GetAuthName_bat(filename,ManualVer).split(" ")[0:3]]
                #Function calls for .ice
                elif(filename.endswith(".ice")):
                        ScriptSynVer=GetSynVer_ice(filename)
                        Test_script_case_count = String_Matching(filename,'Test Procedure Number')
                        ManualVer=GetManVer_ice(filename)
                        SVCPNo=GetSVCPNo(filename)
                        if(type_flag==1):
                                AuthName=GetAuthName_ice_t1(filename)
                        elif(type_flag==2):
                                AuthName=[element+" " for element in GetAuthName_ice(filename,ManualVer).split(" ")[0:3]]
                        else:
                                AuthName=" "
                #Function calls for .docx
                elif(filename.endswith(".docx")):
                        Entered_docx=1
                        SVCPNo,docx_svcp_flag=Docx_GetSVCPNo(filename)
                        Author=Docx_GetAuthor(filename)
                        Test_script_case_count=Docx_GetResultTCcount(filename)

                        
                if ((Entered_docx==0 and docx_svcp_flag==0) or (Entered_docx==1 and docx_svcp_flag==1)):
                    for entry in SVCPNo:
                            TPS=GetTPSID(entry)
                            SWRDandSWDD=GetSWRDandSWDD(entry)
                            sheet1.write(i,0,FileWithoutPath)
                            sheet1.write(i,1,filename.split(FileWithoutPath)[0])
                            if(".docx" not in filename):
                                sheet1.write(i,2,ScriptSynVer)
                                sheet1.write(i,3,ManualVer)
                                if(ScriptSynVer==ManualVer):
                                        sheet1.write(i,4,"Match")
                                else:
                                        sheet1.write(i,4,"Version mismatch")
                            sheet1.write(i,5,AuthName)
                            sheet1.write(i,6,str(entry))
                            if TPS=="ID" or TPS==" ID" or TPS==" ID " or TPS=="ID "\
                               or TPS=="ID\n" or TPS=="\nID" or TPS=="\nID\n" or TPS=="ID\t" :
                                    TPS=" "
                            sheet1.write(i,7,TPS)
                            if SWRDandSWDD=="Out-links at depth 2":
                                SWRDandSWDD=" "
                            sheet1.write(i,8,SWRDandSWDD)
                            

                            ResultTestScriptVersion=GetResultTestScriptVersion(filename_res)
                            ResTestScriptVersion=GetResTestScript(filename_res)

                            
                            if(ResTestScriptVersion.endswith(".java")):
                                result=GetResult_java(filename_res)
                                Test_Result_case_count = Reverse_reading(filename_res, 'TC')
                                tc_id=GetTCid_java(filename_res)
                                result_version=GetResult_version(filename_res)
                                
                            elif(ResTestScriptVersion.endswith(".bat")):
                                result=GetResult_bat(filename_res)
                                Test_Result_case_count = String_Matching(filename_res,'Test Case ID')
                                tc_id=GetTCid_bat(filename_res)
                                result_version=GetResult_version(filename_res)
                                
                            elif(ResTestScriptVersion.find('.ice') != -1):
                                result=GetResult_ice(filename_res)
                                Test_Result_case_count = String_Matching(filename_res,'Test Procedure Number')
                                tc_id=GetTCid_ice(filename_res)
                                result_version=GetResult_version(filename_res)
                                
                            elif(".docx" in filename_res):
                                result=Docx_GetResult(filename_res)
                                result_version=Docx_GetResultVersion(filename_res)
                                Test_Result_case_count=Docx_GetResultTCcount(filename_res)
                                tc_id=Docx_GetResultTCID(filename_res)
                                
                            
                            sheet1.write(i,9,FileWithoutPath_res)
                            sheet1.write(i,10,filename_res.split(FileWithoutPath_res)[0])
                            sheet1.write(i,13,result)
                            if result_version==0:
                                    result_version=""
                            sheet1.write(i,11,result_version)
                            sheet1.write(i,14,tc_id)
                            sheet1.write(i,15,Test_script_case_count)
                            sheet1.write(i,16,Test_Result_case_count)
                            if(Test_script_case_count==Test_Result_case_count):
                                sheet1.write(i,17,"Count Match")
                            else:
                                sheet1.write(i,17,"Count Mismatch")

                            if(".docx" not in filename):
                                
                                
                                
                                sheet1.write(i,12,ResultTestScriptVersion)
                                
                                if(ScriptSynVer==ResultTestScriptVersion):
                                        sheet1.write(i,18,"Match")
                                else:
                                        sheet1.write(i,18,"Version mismatch")
                                match_var=Time_Stamp(filename,filename_res)
                                sheet1.write(i,19,match_var)
                            i=i+1
                else:
                        result=Docx_GetResult(filename_res)
                        SWRDandSWDD=Docx_GetSWRDandSWDD_1(filename)
                        result_version=Docx_GetResultVersion(filename_res)
                        tc_id=Docx_GetResultTCID(filename_res)
                        Test_Result_case_count=Docx_GetResultTCcount(filename_res)
                        sheet1.write(i,0,FileWithoutPath)
                        sheet1.write(i,1,filename.split(FileWithoutPath)[0])                        
                        sheet1.write(i,5,AuthName)
                        sheet1.write(i,8,SWRDandSWDD)
                        sheet1.write(i,9,FileWithoutPath_res)
                        sheet1.write(i,10,filename_res.split(FileWithoutPath_res)[0])
                        sheet1.write(i,11,result_version)
                        sheet1.write(i,13,result)
                        sheet1.write(i,14,tc_id)
                        sheet1.write(i,15,Test_script_case_count)
                        sheet1.write(i,16,Test_Result_case_count)
                        if(Test_script_case_count==Test_Result_case_count):
                             sheet1.write(i,17,"Count Match")
                        else:
                             sheet1.write(i,17,"Count Mismatch")
                        i=i+1
                        
    workbook_name='Inter'+date_time+'.xls'
    wb.save(workbook_name)

    wb_1=in_wb = openpyxl.load_workbook(workbook_name)
    sheet_wb= wb_1.worksheets[0]
    for i in range (1,sheet_wb.max_row):
            file_list.append([sheet_wb.cell(row=i,column=6).value,\
                              sheet_wb.cell(row=i,column=0).value,sheet_wb.cell(row=i,column=9).value,\
                              " "," "," ",sheet_wb.cell(row=i,column=2).value,sheet_wb.cell(row=i,column=11).value])
            if("," in sheet_wb.cell(row=i,column=6).value):
                    svcp=sheet_wb.cell(row=i,column=6).value.split(",")
                    for sv in svcp:
                            svcp_in_files.append(sv)
            else:
                    svcp_in_files.append(sheet_wb.cell(row=i,column=6).value)

            Script_Count.append([sheet_wb.cell(row=i,column=0).value])
            if(".docx" not in sheet_wb.cell(i,0).value):
                    Script_Count_without_docx.append([sheet_wb.cell(row=i,column=0).value])
                    
            Results_Count.append([sheet_wb.cell(row=i,column=9).value])            
            if(".docx" not in sheet_wb.cell(row=i,column=9).value):
                    Results_Count_without_docx.append([sheet_wb.cell(row=i,column=0).value])
                    
            SVCP_Count.append( sheet_wb.cell(row=i,column=6).value)
            if("Version mismatch" in sheet_wb.cell(row=i,column=4).value):
                    Script_Mismatch.append([sheet_wb.cell(row=i,column=0).value])
            if("Count Mismatch" in sheet_wb.cell(row=i,column=17).value ):
                    Test_Case_Mismatch.append([sheet_wb.cell(row=i,column=9).value])
            if("Version mismatch" in sheet_wb.cell(row=i,column=18).value ):
                    Syn_Result_TestScript.append([sheet_wb.cell(row=i,column=9).value])
            if("Time stamp is mismatching" in sheet_wb.cell(row=i,column=19).value ):
                    Time_Stamp_Mismatch.append([sheet_wb.cell(row=i,column=9).value])
            if("ID"==sheet_wb.cell(i,7).value or sheet_wb.cell(row=i,column=7).value == ''):
                    TPS_list.append( sheet_wb.cell(row=i,column=6).value)


    wb_1.release_resources()
    del wb_1



    file_list_with_checklist=Checklist(file_list)

    wb_3= open_workbook(workbook_name)
    wb_2 = copy(wb_3)
    sheet_2=wb_2.add_sheet("Blank SVCP IDs")
    sheet_2.write(0,0,"Script SVCP IDs not found in Excel")
    sheet_2.write(0,1,"Excel SVCP IDs not found in Script")
    r=1
    for f in file_list_with_checklist:
            if(f[3]=="File Not Found"):
                    Not_Found_STC.append(f[0])
            if(f[4]=="File Not Found"):
                    Not_Found_STS.append(f[1])
            if(f[5]=="File Not Found"):
                    Not_Found_STR.append(f[2])


            wb_2.get_sheet(0).write(r,20,f[3])
            wb_2.get_sheet(0).write(r,21,f[4])
            wb_2.get_sheet(0).write(r,22,f[5])
            r=r+1
    r=1
    SVCP_not_in_excel_1=[]
    for s in svcp_in_files:
            if s not in SVCP_List:
                    sheet_2.write(r,0,s)
                    SVCP_not_in_excel_1.append(s)
                    r=r+1
    SVCP_not_in_excel_1=RemoveDuplicate(SVCP_not_in_excel_1)
    SVCP_not_in_excel_c=len(SVCP_not_in_excel_1)

    r=1
    for s in SVCP_List:
            if s not in svcp_in_files:
                    sheet_2.write(r,1,s)
                    r=r+1
    SVCP_not_in_script_c=r-1



    now = datetime.datetime.now()
    date_time=now.strftime("_%m-%d-%Y_%Hh-%Mm-%Ss")
    final_excel_name='Analysis'+date_time+'.xls'
    wb_2.save(final_excel_name)
    os.remove(workbook_name)



    return SVCP_not_in_excel_c,SVCP_not_in_script_c,final_excel_name





###################################################################################



def main():

                start_time = time.time()
                empty = Label(window,text="Execution Status", font=("Times New Roman Bold", 12))
                empty.place(x=10, y=220+50)
                input_frame1 = Frame(window, width = 200, height = 22 , bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
                input_frame1.place(x=130, y=223+50)
                w = Label(input_frame1,text="IN PROGRESS   ", bg="Yellow",font=("Times New Roman", 10)).pack()
                validate_but=Button(window,text=" Validate Inputs ", font=("Times New Roman", 12),state=DISABLED)
                validate_but.place(x=80,y=205)

                DateFormat = "%d%b%Y_%I-%M-%S-%p"
                SrartTimeStamp = datetime.datetime.now()

                global work_path
                global log_fo
                global Synergy_path
                global Final_Revision
                global SVCP_List
                global Script_Mismatch
                global Test_Case_Mismatch
                global Syn_Result_TestScript
                global Time_Stamp_Mismatch
                global Not_Found_STC
                global Not_Found_STS
                global Not_Found_STR
                global TPS_list
                global Script_Count
                global Results_Count
                global Script_Count_without_docx
                global Results_Count_without_docx
                global SVCP_Count
                global SVCP_Version
                global Script_Count_c
                global Results_Count_c
                global Script_Count_without_docx_c
                global Results_Count_without_docx_c
                global SVCP_Count_c

                global Not_Found_STC_c
                global Not_Found_STS_c
                global Not_Found_STR_c
                global Script_Mismatch_c
                global Test_Case_Mismatch_c
                global Syn_Result_TestScript_c
                global Time_Stamp_Mismatch_c
                global TPS_list_c
                global SVCP_not_in_excel
                global SVCP_not_in_script
                global error_log_file
                global SCVP_in_Excel_c
                global date_time1
                global DocxFileList

                DocxFileList=[]

                path=os.getcwd()
                now = datetime.datetime.now()
                date_time1=now.strftime("_%m-%d-%Y_%Hh-%Mm-%Ss")
                error_log_file=open("error_log"+date_time1+".txt","w+")

                Script_Mismatch_c=0
                Test_Case_Mismatch_c=0
                Syn_Result_TestScript_c=0
                Time_Stamp_Mismatch_c=0
                SVCP_not_in_excel=0
                SVCP_not_in_script=0
                Not_Found_STC_c=0
                Not_Found_STS_c=0
                Not_Found_STR_c=0
                TPS_list_c=0

                Not_Found_STC=[]
                Not_Found_STS=[]
                Not_Found_STR=[]
                TPS_list=[]
                Script_Count=[]
                Results_Count=[]
                Script_Count_without_docx=[]
                Results_Count_without_docx=[]
                
                SVCP_Count=[]


                Script_Mismatch=[]
                Test_Case_Mismatch=[]
                Syn_Result_TestScript=[]
                Time_Stamp_Mismatch=[]


                SVCP_List=[]
                Final_Revision=final_version_text_box.get()


                work_path = 'E2SPDA-SynergyDataAnalysisReport-'+SrartTimeStamp.strftime(DateFormat)+'\\'

                print '==> Collection of File Names From Synergy Folder Started ...'
                global MasterFileNameList
                global FileWithoutPath
                global ResultFileList
                Synergy_path = BrowseFolderPath()
                MasterFileNameList = GetListOfFilesFromFolder (Synergy_path, FileTypeList_withoutResult)         # Store all the script name in MasterFileNameList
                ExcelFileList=GetListOfFilesFromFolder (Synergy_path, ExcelFiles)
                ResultFileList=GetListOfFilesFromFolder (Synergy_path, ResultFileTypeList)
                all_docx_list=GetListOfFilesFromFolder (Synergy_path, docx_type)
                for docxfilename in all_docx_list:
                    DocxFileWithoutPath=docxfilename.split("\\")[-1]
                    if (DocxFileWithoutPath.upper().find('RESULT')!= -1):
                        ResultFileList.append(docxfilename)
                    else:
                        DocxFileList.append(docxfilename)
                
                if len(MasterFileNameList)!= 0:
                        GetSVCPList()
                        SVCP_not_in_excel,SVCP_not_in_script,final_excel_name = ConsolidateData(ExcelFileList)

                        Script_Mismatch=RemoveDuplicate(Script_Mismatch)
                        Test_Case_Mismatch=RemoveDuplicate(Test_Case_Mismatch)
                        Syn_Result_TestScript=RemoveDuplicate(Syn_Result_TestScript)
                        Time_Stamp_Mismatch=RemoveDuplicate(Time_Stamp_Mismatch)
                        TPS_list=RemoveDuplicate(TPS_list)
                        Script_Count=RemoveDuplicate(Script_Count)
                        Results_Count=RemoveDuplicate(Results_Count)
                        Script_Count_without_docx=RemoveDuplicate(Script_Count_without_docx)
                        Results_Count_without_docx=RemoveDuplicate(Results_Count_without_docx)
                        SVCP_Count=RemoveDuplicate(SVCP_Count)

                        Not_Found_STC=RemoveDuplicate(Not_Found_STC)
                        Not_Found_STS=RemoveDuplicate(Not_Found_STS)
                        Not_Found_STR=RemoveDuplicate(Not_Found_STR)
                        SVCP_List=RemoveDuplicate(SVCP_List)

                        Script_Mismatch_c=len(Script_Mismatch)
                        Test_Case_Mismatch_c=len(Test_Case_Mismatch)
                        Syn_Result_TestScript_c=len(Syn_Result_TestScript)
                        Time_Stamp_Mismatch_c=len(Time_Stamp_Mismatch)
                        TPS_list_c=len(TPS_list)
                        Script_Count_c=len(Script_Count)
                        Results_Count_c=len(Results_Count)
                        
                        Script_Count_without_docx_c=len(Script_Count_without_docx)
                        Results_Count_without_docx_c=len(Results_Count_without_docx)
                        
                        SVCP_Count_c=len(SVCP_Count)
                        SCVP_in_Excel_c=len(SVCP_List)

                        Not_Found_STC_c=len(Not_Found_STC)
                        Not_Found_STS_c=len(Not_Found_STS)
                        Not_Found_STR_c=len(Not_Found_STR)

                        now = datetime.datetime.now()
                        date_time=now.strftime("_%m-%d-%Y_%Hh-%Mm-%Ss")
                        mismatch_name='Mismatch_Count'+date_time+'.txt'

                        count_mismatch_file=open(mismatch_name,"w+")

                        count_mismatch_file.write(" Mismatches Count Below")
                        count_mismatch_file.write("\nSynergy Version Vs Manual Version = "+str(Script_Mismatch_c)+"/"+str(Script_Count_without_docx_c))
                        count_mismatch_file.write("\nTest case count in script Vs Test case count in Result = "+str(Test_Case_Mismatch_c)+"/"+str(Results_Count_c))
                        count_mismatch_file.write("\nSynergy version mismatch in Result = "+str(Syn_Result_TestScript_c)+"/"+str(Results_Count_without_docx_c))
                        count_mismatch_file.write("\nTime stamp Mismatch TS Vs TR = " +str(Time_Stamp_Mismatch_c)+"/"+str(Results_Count_without_docx_c))
                        count_mismatch_file.write("\nNo. of SVCP not linked to Requirement = "+str(SVCP_not_in_excel)+"/"+str(SVCP_Count_c))
                        count_mismatch_file.write("\nNo. of SVCP not linked to Script = "+str(SVCP_not_in_script)+"/"+str(SCVP_in_Excel_c))
                        count_mismatch_file.write("\nSVCP not found in Review Checklist = "+str(Not_Found_STC_c)+"/"+str(SVCP_Count_c))
                        count_mismatch_file.write("\nScripts not found in Review Checklist = "+str(Not_Found_STS_c)+"/"+str(Script_Count_c))
                        count_mismatch_file.write("\nResults not found in Review Checklist = "+str(Not_Found_STR_c)+"/"+str(Results_Count_c))
                        count_mismatch_file.write("\nNo. of SVCP ID not linked to TPS ID = "+str(TPS_list_c)+"/"+str(SVCP_Count_c))


                        count_mismatch_file.close()

                        EndTimeStamp = datetime.datetime.now()

                        TimeTaken = EndTimeStamp-SrartTimeStamp

                        total_seconds = time.time() - start_time
                        seconds_1 = total_seconds % (24 * 3600)
                        hour = seconds_1 // 3600
                        seconds_1 %= 3600
                        minutes = seconds_1 // 60
                        seconds_1 %= 60
                        run_complete(final_excel_name,hour,minutes,seconds_1)
                        error_log_file.close()
                        validate_but=Button(window,text=" Validate Inputs ", font=("Times New Roman", 12),command=sheet_validate)
                        validate_but.place(x=80,y=205)

                else:
                        print '*******No Files Selected*******'
                        tkMessageBox.showinfo("Error"," No Files are Selected")







def sheet_validate():
    main_flag=2
    main1_flag=2
    main2_flag=2
    global wb_main_app_svcp
    global sheet_mainapp
    global wb_main_app_svcp_1
    global sheet_mainapp_1
    global wb_main_app_svcp_2
    global sheet_mainapp_2
    validate_but=Button(window,text=" Validate Inputs ", font=("Times New Roman", 12),state=DISABLED)
    validate_but.place(x=80,y=205)


    if len(final_version_text_box.get()) == 0:
        tkMessageBox.showinfo("Error"," Enter the SVCP Baseline version")
        validate_but=Button(window,text=" Validate Inputs ", font=("Times New Roman", 12),command=sheet_validate)
        validate_but.place(x=80,y=205)
    else:
        teams = range(6)
        popup = tk.Toplevel()
        tk.Label(popup, width = 40, height = 3,text="Inputs are validating.....").grid(row=0,column=0)
        progress = 0
        progress_var = tk.DoubleVar()
        progress_bar = ttk.Progressbar(popup,variable=progress_var, maximum=100)
        progress_bar.grid(row=2, column=0)#.pack(fill=tk.X, expand=1, side=tk.BOTTOM)
        popup.pack_slaves()
        progress_step = float(100.0/len(teams))
        for team in teams:
           popup.update()
           sleep(4) # launch task
           progress += progress_step
           progress_var.set(progress)

        #TPS
        try:
           wb_main_app_svcp=openpyxl.load_workbook(tpspath)
           sheet_mainapp = wb_main_app_svcp.worksheets[0]
           flag_1=0
           flag_2=0
           flag_3=0
           for col in range(1,sheet_mainapp.max_column+1):
               if("ID" in str(sheet_mainapp.cell(row=1,column=col).value)):
                              flag_1=1
               if("Out-links at depth 1" in str(sheet_mainapp.cell(row=1,column=col).value)):
                              flag_2=2
##               if("Out-links at depth" in sheet_mainapp.cell_value(0,col)):
##                              flag_3=3

           if(flag_1==0 and flag_2==0):
                main_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) ID and Out-links at depth 1 are missing in the TPS_SVCP file\n Please select the correct file")
           elif(flag_1==0 ):
                main_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) ID is missing in the TPS_SVCP file\n Please select the correct file")
           elif(flag_2==0):
                main_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) Out-links at depth 1 is missing in the TPS_SVCP file\n Please select the correct file")
           else:
                main_flag=0

           if(main_flag==1):
                input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
                input_frame.place(x=180,y=130)
                lbl2=Label(input_frame, text= "                                                  " , fg='Black', width = 30, height = 1,font=("Times New Roman", 10),bg="White").pack()

        except Exception as e:
           print("TCP",e)
           tkMessageBox.showinfo("Error", "TPS_SVCP_TR.xlsx file is not selected\n")
           tk.Label(popup, width = 40, height = 3,text="Inputs are not proper").grid(row=0,column=0)

        #Main App
        try:
           wb_main_app_svcp_1=openpyxl.load_workbook(swddpath)
           sheet_mainapp_1 = wb_main_app_svcp_1.worksheets[0]
           flag_1=0
           flag_2=0
           flag_3=0
           for col in range(1,sheet_mainapp_1.max_column+1):
               if("ID" in str(sheet_mainapp_1.cell(row=1,column=col).value)):
                              flag_1=1
               if("Out-links at depth 1" in str(sheet_mainapp_1.cell(row=1,column=col).value)):
                              flag_2=2
               if("DS10793/327" in str(sheet_mainapp_1.cell(row=1,column=col).value)):
                              flag_3=3

           if(flag_1==0 and flag_2==0 and flag_3==0):
                main1_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) ID, Out-links at depth 1 and DS10793/327 are missing in the Main App file\n Please select the correct file")
           elif(flag_1==0 ):
                main1_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) ID is missing in the Main App file\n Please select the correct file")
           elif(flag_2==0 ):
                main1_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) Out-links at depth 1 is missing in the Main App file\n Please select the correct file")
           elif( flag_3==0):
                main1_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) DS10793/327 is missing in the Main App file\n Please select the correct file")
           else:
                main1_flag=0
           if( main1_flag==1):
                input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
                input_frame.place(x=180,y=70)
                lbl2=Label(input_frame, text="                                                  " , width = 30, height = 1,fg='Black', font=("Times New Roman", 10),bg="White").pack()

        except:
           tkMessageBox.showinfo("Error", "Main App file.xlsx file is not selected\n")
           tk.Label(popup, width = 40, height = 3,text="Inputs are not proper").grid(row=0,column=0)


        #MLCT
        try:
           wb_main_app_svcp_2=openpyxl.load_workbook(mlctpath)
           sheet_mainapp_2 = wb_main_app_svcp_2.worksheets[0]
           flag_1=0
           flag_2=0
           flag_3=0
           for col in range(1,sheet_mainapp_2.max_column+1):
               if("Out-links at depth 2" in str(sheet_mainapp_2.cell(row=1,column=col).value)):
                              flag_1=1


               if("Out-links at depth 1" in str(sheet_mainapp_2.cell(row=1,column=col).value)):
                              flag_2=2

               if("DS10793/327 - E2 SPDA Software Verification Cases and Procedures (MLCT is currently part of E2 SPDA SVCP)" in str(sheet_mainapp_2.cell(row=1,column=col).value)):
                              flag_3=3
           if(flag_1==0 and flag_2==0 and flag_3 ==0):
                main2_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) Out-links at depth 2, Out-links at depth 1 and DS10793/327 are missing in the MLCT_SVCP file\n Please select the correct file")
           elif(flag_1==0):
                main2_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) Out-links at depth 2 is missing in the MLCT_SVCP file\n Please select the correct file")
           elif(flag_2==0):
                main2_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) Out-links at depth 1 is missing in the MLCT_SVCP file\n Please select the correct file")
           elif(flag_3 ==0):
                main2_flag=1
                tkMessageBox.showinfo("Error"," Attribute(s) DS10793/327 is missing in the MLCT_SVCP file\n Please select the correct file")
           else:
                main2_flag=0
           if main2_flag==1:
                   input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
                   input_frame.place(x=180,y=100)
                   lbl2=Label(input_frame, text= "                                                  " , fg='Black', width = 30, height = 1,font=("Times New Roman", 10),bg="White").pack()

        except Exception as e:
	   print(e)
           tkMessageBox.showinfo("Error", "MLCT_SVCP file.xlsx file is not selected\n")
           tk.Label(popup, width = 40, height = 3,text="Inputs are not proper").grid(row=0,column=0)
        if(main_flag==0 and main1_flag==0 and main2_flag==0 ):
                tkMessageBox.showinfo("   ","All the inputs are valid. Proceed to Run")
                run_but['state'] = 'normal'
                abort_but['state'] = 'normal'
                tk.Label(popup, width = 40, height = 3,text="Inputs are validated").grid(row=0,column=0)
        validate_but=Button(window,text=" Validate Inputs ", font=("Times New Roman", 12),command=sheet_validate)
        validate_but.place(x=80,y=205)






def run_complete(final_excel_name,h,m,s):







        input_text = StringVar()

        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
        input_frame.place(x=850, y=70+30) #1
        if TPS_list_c==0:
              l11 = Label (input_frame, width = 12, height = 1,text=" "+ str(TPS_list_c)+"/"+str(SVCP_Count_c) +" ", bg="Lightgreen",font=("Times New Roman", 11)).pack()
        else:
              l11 = Label (input_frame, width = 12, height = 1,text=" "+ str(TPS_list_c)+"/"+str(SVCP_Count_c) +" ", bg="lightcoral",font=("Times New Roman", 11)).pack()

        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
        input_frame.place(x=850, y=110+20)#2
        if SVCP_not_in_excel==0:
            l12 = Label (input_frame,width = 12, height = 1, text=" "+ str(SVCP_not_in_excel)+"/"+str(SVCP_Count_c) +" ", bg="Lightgreen",font=("Times New Roman", 11)).pack()
        else:
            l12 = Label (input_frame,width = 12, height = 1, text=" "+ str(SVCP_not_in_excel)+"/"+str(SVCP_Count_c) +" ", bg="lightcoral",font=("Times New Roman", 11)).pack()


        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "White", highlightthickness = 2)
        input_frame.place(x=850, y=170+30)#3
        if Script_Mismatch_c==0:
           l13 = Label (input_frame,width = 12, height = 1, text= " "+ str(Script_Mismatch_c)+"/"+str(Script_Count_without_docx_c) +" ", bg="Lightgreen",font=("Times New Roman", 11)).pack()
        else:
           l13 = Label (input_frame,width = 12, height = 1, text= " "+ str(Script_Mismatch_c)+"/"+str(Script_Count_without_docx_c) +" ", bg="lightcoral",font=("Times New Roman", 11)).pack()

        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
        input_frame.place(x=850, y=170+60)#4
        if SVCP_not_in_script==0:
            l14 = Label (input_frame, width = 12, height = 1, text=" "+ str(SVCP_not_in_script)+"/"+str(SCVP_in_Excel_c) +" ", bg="Lightgreen",font=("Times New Roman", 11)).pack()
        else:
            l14 = Label (input_frame, width = 12, height = 1, text=" "+ str(SVCP_not_in_script)+"/"+str(SCVP_in_Excel_c) +" ", bg="lightcoral",font=("Times New Roman", 11)).pack()


        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
        input_frame.place(x=850, y=270+30)#5
        if Test_Case_Mismatch_c==0:
          l15 = Label (input_frame, width = 12, height = 1,text=" "+ str(Test_Case_Mismatch_c)+"/"+str(Results_Count_c) +" ",bg="Lightgreen", font=("Times New Roman", 11)).pack()
        else:
          l15 = Label (input_frame, width = 12, height = 1,text=" "+ str(Test_Case_Mismatch_c)+"/"+str(Results_Count_c) +" ",bg="lightcoral", font=("Times New Roman", 11)).pack()



        input_frame = Frame(window, width = 200, height = 22 , bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
        input_frame.place(x=850, y=270+60)#6
        if Syn_Result_TestScript_c==0:
           l16 = Label (input_frame,width = 12, height = 1, text=" "+ str(Syn_Result_TestScript_c)+"/"+str(Results_Count_without_docx_c) +" ",bg="Lightgreen", font=("Times New Roman", 11)).pack()
        else:
           l16 = Label (input_frame,width = 12, height = 1, text=" "+ str(Syn_Result_TestScript_c)+"/"+str(Results_Count_without_docx_c) +" ",bg="lightcoral", font=("Times New Roman", 11)).pack()

        input_frame = Frame(window, width = 200, height = 22 , bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
        input_frame.place(x=850, y=270+90)#7
        if Syn_Result_TestScript_c==0:
           l17 = Label (input_frame,width = 12, height = 1, text=" "+ str(Time_Stamp_Mismatch_c)+"/"+str(Results_Count_without_docx_c) +" ", bg="Lightgreen",font=("Times New Roman", 11)).pack()
        else:
           l17 = Label (input_frame,width = 12, height = 1, text=" "+ str(Time_Stamp_Mismatch_c)+"/"+str(Results_Count_without_docx_c) +" ", bg="lightcoral",font=("Times New Roman", 11)).pack()


        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
        input_frame.place(x=850, y=410+30)#8
        if Not_Found_STC_c==0:
            l18 = Label (input_frame,width = 12, height = 1, text=" "+ str(Not_Found_STC_c)+"/"+str(SVCP_Count_c) +" ", bg="Lightgreen",font=("Times New Roman", 11)).pack()
        else:
            l18 = Label (input_frame, width = 12, height = 1, text=" "+ str(Not_Found_STC_c)+"/"+str(SVCP_Count_c) +" ", bg="lightcoral",font=("Times New Roman", 11)).pack()

        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
        input_frame.place(x=850, y=410+60) #9
        if Not_Found_STS_c==0:
              l19 = Label (input_frame, width = 12, height = 1,text=" "+ str(Not_Found_STS_c)+"/"+str(Script_Count_c) +" ",bg="Lightgreen", font=("Times New Roman", 11)).pack()
        else:
              l19 = Label (input_frame, width = 12, height = 1,text=" "+ str(Not_Found_STS_c)+"/"+str(Script_Count_c) +" ",bg="lightcoral", font=("Times New Roman", 11)).pack()

        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
        input_frame.place(x=850, y=410+90) #10
        if Not_Found_STR_c==0:
              l20 = Label (input_frame,width = 12, height = 1, text=" "+ str(Not_Found_STR_c)+"/"+str(Results_Count_c) +" ",bg="Lightgreen", font=("Times New Roman", 11)).pack()
        else:
              l20 = Label (input_frame,width = 12, height = 1, text=" "+ str(Not_Found_STR_c)+"/"+str(Results_Count_c) +" ",bg="lightcoral", font=("Times New Roman", 11)).pack()


        empty = Label(window,text="Execution Status", font=("Times New Roman Bold", 12))
        empty.place(x=10, y=220+50)
        input_frame1 = Frame(window, width = 200, height = 22 , bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
        input_frame1.place(x=130, y=223+50)
        w = Label(input_frame1,text="COMPLETED    ", bg="lightskyblue", font=("Times New Roman", 10)).pack()


        empty = Label(window,text="Elapsed time", font=("Times New Roman Bold", 12))
        empty.place(x=10, y=270+50)
        input_frame1 = Frame(window, width = 200, height = 22 , bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
        input_frame1.place(x=130, y=270+50)
        w = Label(input_frame1,text="%02d:%02d:%02d" %(h, m, s), font=("Times New Roman", 10)).pack()


        path=os.getcwd()
        Excel_filepath=path+"\\"+final_excel_name
        #date_time=""
        error_log=path+"\\error_log"+date_time1+".txt"

        def openerror():
            os.startfile(error_log, 'open')
        def openexcel():
            os.startfile(Excel_filepath, 'open')

        lbl1=Label(window, text="Error Log Path:", fg='Black', font=("Times New Roman bold", 10))
        lbl1.place(x=10,y=560)
        lbl2=Label(window, text=error_log, fg='Black', bd=2,font=("Times New Roman", 10))
        lbl2.place(x=10,y=580)
        lbl3=Label(window, text="Output Report Path:", fg='Black', font=("Times New Roman bold", 10))
        lbl3.place(x=10,y=600)
        lbl4=Label(window, text=Excel_filepath, fg='Black',bd=2, font=("Times New Roman", 10))
        lbl4.place(x=10,y=620)
        button2 = Button(window, text="  Open  ", command=openerror)
        button2.place(x=700,y=570)
        button3 = Button(window, text="  Open  ", command=openexcel)
        button3.place(x=700,y=620)




def stop_fun():
        window.destroy()
        sys.exit()

def browse_swdd():
        global swddpath
        swddpath = tkFileDialog.askopenfilename(filetypes=(("Template files","*.xlsx"),("All files","*.xlsx")))
        disp_path=swddpath.split("/")[-1]
        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
        input_frame.place(x=180,y=70)
        lbl2=Label(input_frame, text= disp_path , fg='Black', font=("Times New Roman", 10),width = 30, height = 1,bg="White").pack()
        print(swddpath)

def browse_mlct():
        global mlctpath
        mlctpath = tkFileDialog.askopenfilename(filetypes=(("Template files","*.xlsx"),("All files","*.xlsx")))
        disp_path=mlctpath.split("/")[-1]
        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
        input_frame.place(x=180,y=100)
        lbl2=Label(input_frame, text= disp_path , fg='Black', font=("Times New Roman", 10),width = 30, height = 1,bg="White").pack()
        print(mlctpath)

def browse_tps():
     try:
        global tpspath
        tpspath = tkFileDialog.askopenfilename(filetypes=(("Template files","*.xlsx"),("All files","*.xlsx")))
        disp_path=tpspath.split("/")[-1]
        input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
        input_frame.place(x=180,y=130)
        lbl2=Label(input_frame, text= disp_path , fg='Black', font=("Times New Roman", 10),width = 30, height = 1,bg="White").pack()
        print(tpspath)
     except:
         tkMessageBox.showinfo("TPS_SVCP_TR.xlsx file is not selected\n")

def Help():
    from os.path import exists
    window1 = Tkinter.Tk()
    window1.geometry("650x630")
    window1.title("Instructions" )

    lbl=Label(window1, text="Steps to execute the E2 SPDA Synergy Analysis Tool:", fg='Black', font=("Times New Roman Bold", 16),anchor=CENTER,bg="light blue", padx=1000)
    lbl.pack(side="top")

##    lbl2=Label(window1, text="Steps to execute the E2 SPDA Synergy Analysis Tool:",bd=2,fg='Black',anchor=CENTER,bg="light blue", font=("Arial Bold", 15))
##    lbl2.place(x=50,y=10)
    lbl3=Label(window1, text="1. Extract Main_App_SVCP from E2 SPDA doors with the following attribute.",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=60)
    lbl3=Label(window1, text="     - ID",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=60+30)
    lbl3=Label(window1, text="     - DS10793/327",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=110)
    lbl3=Label(window1, text="     - In-links at depth 1",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=130)
    lbl3=Label(window1, text="     - Out-links at depth 1",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=150)
    lbl3=Label(window1, text="     - Out-links at depth 2",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=170)


    lbl3=Label(window1, text="2. Extract MCLT_SVCP from E2 SPDA doors with the following attribute",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=200)
    lbl3=Label(window1, text="     - ID",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=220)
    lbl3=Label(window1, text="     - DS10793/327 - E2 SPDA Software Verification Cases and Procedures",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=240)
    lbl3=Label(window1, text="     - In-links at depth 1",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=260)
    lbl3=Label(window1, text="     - In-links at depth 2",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=280)
    lbl3=Label(window1, text="     - Out-links at depth 1",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=300)
    lbl3=Label(window1, text="     - Out-links at depth 2",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=320)


    lbl3=Label(window1, text="3. Extract TPS SVCP surrogate from E2 SPDA doors with following attribute.",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=350)
    lbl3=Label(window1, text="     - ID",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=370)
    lbl3=Label(window1, text="     - In-links at depth 1",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=390)
    lbl3=Label(window1, text="     - Out-links at depth",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=410)


    lbl3=Label(window1, text="4. Select 3 input files in E2 SPDA GUI.",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=440)
    lbl3=Label(window1, text="5. Enter the Main app SVCP baseline version",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=460)
    lbl3=Label(window1, text="6. Click Validate Inputs button.",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=480)
    lbl3=Label(window1, text="7. Once validation passed 'Run' button will be enabled.",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=500)
    lbl3=Label(window1, text="8. Click 'run' button and give the synergy dump path.",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=520)
    lbl3=Label(window1, text="9. Wait for the execution to complete.",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=540)
    lbl3=Label(window1, text="10.'error log' & 'Output file' path will be provided in the GUI bottom.",bd=2,fg='Black',font=("Times New Roman", 11))
    lbl3.place(x=10,y=560)

#GUI
from os.path import exists

window = Tkinter.Tk()
window.geometry("990x680")
window.resizable(0, 0)
window.title("Synergy Analysis Tool - E2 SPDA" )
ratio = 16
horizontal_edge_len = 1
vertical_edge_len = 1
hypotenuse = tk.Frame(window, bg='Black', height=1, width=480)
hypotenuse.place(x=490, y=85)

hypotenuse = tk.Frame(window, bg='Black', height=460, width=1)
hypotenuse.place(x=490, y=85)

hypotenuse = tk.Frame(window, bg='Black', height=460, width=1)
hypotenuse.place(x=970, y=85)

hypotenuse = tk.Frame(window, bg='Black', height=1, width=480)
hypotenuse.place(x=490, y=285)

hypotenuse = tk.Frame(window, bg='Black', height=1, width=480)
hypotenuse.place(x=490, y=185)

hypotenuse = tk.Frame(window, bg='Black', height=1, width=480)
hypotenuse.place(x=490, y=425)

hypotenuse = tk.Frame(window, bg='Black', height=1, width=480)
hypotenuse.place(x=490, y=545)




cwd = os.getcwd()
f_path1 = cwd +'\Capture_Logo.GIF'
if exists(f_path1):
    img = PhotoImage(file = cwd +'\Capture_Logo.GIF' )
    window.tk.call('wm', 'iconphoto', window._w, img)


lbl=Label(window, text="Synergy Analysis Tool - E2 SPDA", fg='Black', font=("Times New Roman Bold", 16),anchor=CENTER,bg="light blue", padx=1000)
lbl.pack(side="top")


run_but=Button(window, width = 8, height = int(0.6), text="  Readme  ", bd=2,fg='Black',bg="light blue", font=("Times New Roman Bold", 12), command=Help)
run_but.place(x=850, y=0)


run_but=Button(window, text="        Run       ", bd=2,fg='Black',font=("Times New Roman Bold", 12), state=DISABLED, command=main)
run_but.place(x=10, y=350+30)

abort_but=Button(window, text="       Abort      ", bd=2,fg='Black',font=("Times New Roman Bold", 12),state=DISABLED, command=stop_fun)
abort_but.place(x=10, y=400+30)


stop_but=Button(window, text="  Select Mainapp file          ", bd=2,fg='Black',font=("Times New Roman Bold", 10),command=browse_swdd)
stop_but.place(x=20, y=70)

input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
input_frame.place(x=180,y=70)
lbl2=Label(input_frame, text="                                                          " , width = 30, height = 1, fg='Black', font=("Times New Roman", 10),bg="White").pack()


stop_but=Button(window, text="  Select MLCT_SVCP file", bd=2,fg='Black',font=("Times New Roman Bold", 10),command=browse_mlct)
stop_but.place(x=20, y=100)

input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
input_frame.place(x=180,y=100)
lbl2=Label(input_frame, text= "                                                          ", width = 30, height = 1 , fg='Black', font=("Times New Roman", 10),bg="White").pack()

stop_but=Button(window, text="  Select TPS_SVCP file    ", bd=2,fg='Black',font=("Times New Roman Bold", 10),command=browse_tps)
stop_but.place(x=20, y=130)

input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
input_frame.place(x=180,y=130)
lbl2=Label(input_frame, text= "                                                          ", width = 30, height = 1 , fg='Black', font=("Times New Roman", 10),bg="White").pack()


final_version_label = Label(window, text="Enter SVCP Baseline version",font=("Times New Roman bold", 10))
final_version_label.place(x=10, y=170)

final_version_text_box = Entry(window, bd=4)
final_version_text_box.place(x=180, y=170)

SVCP_Version = final_version_text_box
validate_but=Button(window,text=" Validate Inputs ", font=("Times New Roman", 12),command=sheet_validate)
validate_but.place(x=80,y=205)



l1 = Label (window, text="Number of Mismatches", font=("Times New Roman Bold", 15))
l1.place(x=600, y=30)

l1 = Label (window, text="Number of SVCP Mismatches", font=("Times New Roman Bold", 12))
l1.place(x=500, y=70)



l10 = Label (window, text="1. No. of SVCP ID not linked to TPS ID", font=("Times New Roman", 11))
l10.place(x=500, y=70+30)

l5 = Label (window, text="2. No. of SVCP not linked to Requirement", font=("Times New Roman", 11))
l5.place(x=500, y=110+20)


l1 = Label (window, text="Number of Script Mismatches", font=("Times New Roman Bold", 12))
l1.place(x=500, y=170)


l1 = Label (window, text="3. TS Synergy and Manual Version Mismatch", font=("Times New Roman", 11))
l1.place(x=500, y=170+30)

l6 = Label (window, text="4. No. of SVCP not linked to Script", font=("Times New Roman", 11))
l6.place(x=500, y=170+60)


l1 = Label (window, text="Number of Test Result Mismatches", font=("Times New Roman Bold", 12))
l1.place(x=500, y=270)


l2 = Label (window, text="5. TC count Mismatch between TS and TR", font=("Times New Roman", 11))
l2.place(x=500, y=270+30)


l3 = Label (window, text="6. No. of Synergy version mismatch between TS and TR", font=("Times New Roman", 11))
l3.place(x=500, y=270+60)

l4 = Label (window, text="7. No. of Time stamp Mismatch between TS and TR", font=("Times New Roman", 10))
l4.place(x=500, y=270+90)


l1 = Label (window, text="Number of Review checklists Mismatches", font=("Times New Roman Bold", 12))
l1.place(x=500, y=410)

l7 = Label (window, text="8. No. of SVCP not found in Review Checklist", font=("Times New Roman", 11))
l7.place(x=500, y=410+30)

l8 = Label (window, text="9. No of Scripts not found in Review Checklist", font=("Times New Roman", 11))
l8.place(x=500, y=410+60)

l9 = Label (window, text="10. No. of Results not found in Review Checklist", font=("Times New Roman", 11))
l9.place(x=500, y=410+90)


input_text = StringVar()
input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "White", highlightthickness = 2)
input_frame.place(x=850, y=70+30)
l11 = Label (input_frame,width = 12, height = 1, text= " 0/0 ", bg="white",font=("Times New Roman", 10)).pack()


input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
input_frame.place(x=850, y=110+20)
l12 = Label (input_frame, width = 12, height = 1,text=" 0/0 ",bg="white", font=("Times New Roman", 10)).pack()



input_frame = Frame(window, width = 200, height = 22 , bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
l13 = Label (input_frame,width = 12, height = 1, text=" 0/0 ", bg="white",font=("Times New Roman", 10)).pack()
input_frame.place(x=850, y=170+30)

input_frame = Frame(window, width = 200, height = 22 , bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
input_frame.place(x=850, y=170+60)#
l14 = Label (input_frame,width = 12, height = 1, text=" 0/0 ", bg="white",font=("Times New Roman", 10)).pack()

input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
input_frame.place(x=850, y=270+30)#
l15 = Label (input_frame,width = 12, height = 1, text=" 0/0 ", bg="white",font=("Times New Roman", 10)).pack()

input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
input_frame.place(x=850, y=270+60)
l16 = Label (input_frame, width = 12, height = 1, text=" 0/0", bg="white",font=("Times New Roman", 10)).pack()

input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
input_frame.place(x=850, y=270+90)#
l17 = Label (input_frame,width = 12, height = 1, text=" 0/0 ", bg="white",font=("Times New Roman", 10)).pack()

input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
input_frame.place(x=850, y=410+30) #
l18 = Label (input_frame, width = 12, height = 1,text=" 0/0 ",bg="white", font=("Times New Roman", 10)).pack()

input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
input_frame.place(x=850, y=410+60) #
l19 = Label (input_frame,width = 12, height = 1, text=" 0/0 ",bg="white", font=("Times New Roman", 10)).pack()

input_frame = Frame(window, width = 200, height = 22, bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 2)
input_frame.place(x=850, y=410+90)#
l20 = Label (input_frame, width = 12, height = 1,text=" 0/0 ", bg="white", font=("Times New Roman", 10)).pack()


empty = Label(window,text="Execution Status", font=("Times New Roman Bold", 12))
empty.place(x=10, y=220+50)
input_frame1 = Frame(window, width = 200, height = 22 , bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
input_frame1.place(x=130, y=223+50)
w = Label(input_frame1,text="NOT STARTED", font=("Times New Roman", 10)).pack()


empty = Label(window,text="Elapsed time", font=("Times New Roman Bold", 12))
empty.place(x=10, y=270+50)
input_frame1 = Frame(window, width = 200, height = 22 , bg="White", highlightbackground = "black", highlightcolor = "black", highlightthickness = 1)
input_frame1.place(x=130, y=270+50)
w = Label(input_frame1,text="00:00:00" , font=("Times New Roman", 10)).pack()


lbl2=Label(window, text="Copyright (c) 2020 L&T Technology Services. All Rights Reserved.",fg='Black', font=("Times New Roman", 12), bg="light blue",padx=1000)
lbl2.pack(side="bottom")

##img = ImageTk.PhotoImage(Image.open("logo_ltts.gif"))
##panel = Label(window, image = img, bg="white", padx=1000)
##panel.pack(side = "bottom")


window.mainloop()





